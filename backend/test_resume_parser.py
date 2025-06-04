"""
Test script for resume parsing functionality.
"""
import os
import tempfile
from pathlib import Path

from app.resume_parser import ResumeParser


def create_test_docx():
    """Create a simple test DOCX file for testing."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('张三', 0)
        doc.add_paragraph('邮箱: zhangsan@example.com')
        doc.add_paragraph('电话: 13812345678')
        
        doc.add_heading('教育背景', level=1)
        doc.add_paragraph('2018-2022 北京大学 计算机科学与技术 本科')
        
        doc.add_heading('工作经历', level=1)
        doc.add_paragraph('2022-2024 腾讯科技 软件工程师')
        doc.add_paragraph('负责后端开发和系统维护')
        
        doc.add_heading('技能', level=1)
        doc.add_paragraph('Python, Java, MySQL, Redis')
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name
        
    except ImportError:
        print("python-docx not installed, skipping DOCX test")
        return None


def create_test_pdf():
    """Create a simple test PDF file for testing."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        c = canvas.Canvas(temp_file.name, pagesize=letter)
        c.drawString(100, 750, "张三")
        c.drawString(100, 730, "邮箱: zhangsan@example.com")
        c.drawString(100, 710, "电话: 13812345678")
        c.drawString(100, 680, "教育背景:")
        c.drawString(120, 660, "2018-2022 北京大学 计算机科学与技术 本科")
        c.drawString(100, 630, "工作经历:")
        c.drawString(120, 610, "2022-2024 腾讯科技 软件工程师")
        c.drawString(100, 580, "技能:")
        c.drawString(120, 560, "Python, Java, MySQL, Redis")
        c.save()
        
        return temp_file.name
        
    except ImportError:
        print("reportlab not installed, creating simple text file as PDF substitute")
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        temp_file.write("""张三
邮箱: zhangsan@example.com
电话: 13812345678

教育背景:
2018-2022 北京大学 计算机科学与技术 本科

工作经历:
2022-2024 腾讯科技 软件工程师
负责后端开发和系统维护

技能:
Python, Java, MySQL, Redis
""")
        temp_file.close()
        return temp_file.name


def test_resume_parser():
    """Test the resume parser with sample files."""
    parser = ResumeParser()
    
    print("Testing Resume Parser...")
    print("=" * 50)
    
    docx_file = create_test_docx()
    if docx_file:
        try:
            print("\n1. Testing DOCX parsing:")
            result = parser.parse_file(docx_file)
            print(f"File type: {result['file_type']}")
            print(f"Raw text length: {len(result['raw_text'])}")
            print(f"Name extracted: {result['parsed_data']['name']}")
            print(f"Email extracted: {result['parsed_data']['email']}")
            print(f"Phone extracted: {result['parsed_data']['phone']}")
            print(f"Education items: {len(result['parsed_data']['education'])}")
            print(f"Experience items: {len(result['parsed_data']['experience'])}")
            print(f"Skills items: {len(result['parsed_data']['skills'])}")
            
        except Exception as e:
            print(f"DOCX parsing failed: {e}")
        finally:
            if os.path.exists(docx_file):
                os.unlink(docx_file)
    
    pdf_file = create_test_pdf()
    if pdf_file:
        try:
            print("\n2. Testing PDF/text parsing:")
            if pdf_file.endswith('.txt'):
                print("(Using text file as PDF substitute)")
                with open(pdf_file, 'r', encoding='utf-8') as f:
                    text = f.read()
                parsed_data = parser._extract_structured_data(text)
                result = {
                    'file_type': '.txt',
                    'raw_text': text,
                    'parsed_data': parsed_data
                }
            else:
                result = parser.parse_file(pdf_file)
            
            print(f"File type: {result['file_type']}")
            print(f"Raw text length: {len(result['raw_text'])}")
            print(f"Name extracted: {result['parsed_data']['name']}")
            print(f"Email extracted: {result['parsed_data']['email']}")
            print(f"Phone extracted: {result['parsed_data']['phone']}")
            print(f"Education items: {len(result['parsed_data']['education'])}")
            print(f"Experience items: {len(result['parsed_data']['experience'])}")
            print(f"Skills items: {len(result['parsed_data']['skills'])}")
            
        except Exception as e:
            print(f"PDF/text parsing failed: {e}")
        finally:
            if os.path.exists(pdf_file):
                os.unlink(pdf_file)
    
    print("\n" + "=" * 50)
    print("Resume parser testing completed!")


if __name__ == "__main__":
    test_resume_parser()
