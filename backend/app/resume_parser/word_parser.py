"""
Word document parser for extracting text from .doc and .docx files.
"""
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class WordParser:
    """Parser for Word documents (.doc and .docx files)."""
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Extracted text content
            
        Raises:
            ImportError: If required libraries are not installed
            Exception: If file cannot be read
        """
        try:
            if file_path.lower().endswith('.docx'):
                return self._extract_from_docx(file_path)
            else:
                return self._extract_from_doc(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {e}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from .docx file using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx library is required for .docx files. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading .docx file {file_path}: {e}")
            raise
    
    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from .doc file using docx2txt or alternative methods."""
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text.strip()
        except ImportError:
            logger.warning("docx2txt not available, trying alternative methods")
        except Exception as e:
            logger.warning(f"docx2txt failed for {file_path}: {e}")
        
        try:
            return self._extract_from_docx(file_path)
        except Exception as e:
            logger.warning(f"python-docx failed for .doc file {file_path}: {e}")
        
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            return text.strip() if text else ""
        except ImportError:
            raise ImportError(
                "No suitable library found for .doc files. "
                "Install one of: pip install docx2txt, pip install textract, or pip install python-docx"
            )
        except Exception as e:
            logger.error(f"All methods failed for .doc file {file_path}: {e}")
            raise
